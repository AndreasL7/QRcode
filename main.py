import os
import sys
import platform

import qrcode
from loguru import logger
import hydra
from omegaconf import DictConfig
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Cm, Inches

def get_font(size):
    """Helper to find a valid system font."""
    paths = []
    if platform.system() == "Windows":
        paths = ["C:\\Windows\\Fonts\\arial.ttf", "C:\\Windows\\Fonts\\segoeui.ttf"]
    elif platform.system() == "Darwin":
        paths = ["/Library/Fonts/Arial.ttf", "/System/Library/Fonts/Helvetica.ttc"]
    else:
        paths = ["/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"]
    
    for path in paths:
        try:
            return ImageFont.truetype(path, size)
        except OSError:
            continue
    return ImageFont.load_default()

@hydra.main(version_base=None, config_path="conf", config_name="config")
def generate_batch(cfg: DictConfig):
    # Load settings from config
    font_size = cfg.settings.font_size
    padding = cfg.settings.padding
    box_size = cfg.settings.box_size
    output_dir = hydra.utils.get_original_cwd() + "/" + cfg.settings.output_dir
    
    os.makedirs(output_dir, exist_ok=True)
    font = get_font(font_size)
    
    logger.remove()

    logger.add(
        sys.stdout,
    colorize=True,
        format="{time:YYYY-MM-DD HH:mm:ss} | {level} | "
    "{module}:{function}:{line} - {message}",
        level="INFO"
    )

    logger.info(f"Starting batch generation for {len(cfg.locations)} locations...")

    for data in cfg.locations:
        # 1. Generate QR Code
        qr = qrcode.QRCode(box_size=box_size, border=4)
        qr.add_data(data)
        qr.make(fit=True)
        qr_img = qr.make_image(fill_color="black", back_color="white").convert("RGBA")
        qr_w, qr_h = qr_img.size

        # 2. Measure Text
        bbox = font.getbbox(data)
        text_w = bbox[2] - bbox[0]
        text_h = bbox[3] - bbox[1]

        # 3. Calculate Canvas
        canvas_width = max(qr_w, text_w) + (padding * 2)
        canvas_height = qr_h + text_h + (padding * 3)

        # 4. Create and Assemble
        canvas = Image.new("RGBA", (canvas_width, canvas_height), "white")
        qr_x = (canvas_width - qr_w) // 2
        canvas.paste(qr_img, (qr_x, padding))

        draw = ImageDraw.Draw(canvas)
        text_x = (canvas_width - text_w) // 2
        text_y = qr_h + (padding * 2) - bbox[1]
        draw.text((text_x, text_y), data, fill="black", font=font)

        # 5. Save
        # Sanitize filename (replace slashes/dots if they exist in location names)
        safe_filename = data.replace("/", "-").replace("\\", "-")
        final_path = os.path.join(output_dir, f"{safe_filename}.png")
        canvas.convert("RGB").save(final_path)
        
        logger.info(f"  [✔] Generated: {final_path}")


def create_full_page_qr_doc(output_dir: str, docx_path: str = "qr_codes_landscape_a4.docx"):
    doc = Document()

    # ── Set A4 landscape on the first section (applies to whole doc unless we add sections)
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width  = Cm(29.7)   # A4 landscape
    section.page_height = Cm(21.0)

    # Optional: reduce margins so QR can be larger (default is ~2.54 cm)
    section.left_margin   = Cm(1.0)
    section.right_margin  = Cm(1.0)
    section.top_margin    = Cm(1.0)
    section.bottom_margin = Cm(1.5)  # a bit more at bottom if adding text

    png_files = sorted(
        [f for f in os.listdir(output_dir) if f.lower().endswith(".png")]
    )

    for i, filename in enumerate(png_files):
        # if i > 0:
        #     doc.add_page_break()   # new page for each QR after the first

        full_path = os.path.join(output_dir, filename)
        location_name = os.path.splitext(filename)[0]

        # Add picture – centered by using its own paragraph
        p = doc.add_paragraph()
        p.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER

        run = p.add_run()
        # Try to fill most of the usable height (portrait height is now the limiting dimension)
        # ≈ 21 cm page height − top/bottom margins − some breathing room
        run.add_picture(full_path, height=Cm(18.0))   # ← tune this: 17–19 cm usually fits well

    doc.save(docx_path)
    print(f"Saved: {docx_path}  ({len(png_files)} pages)")

if __name__ == "__main__":
    generate_batch()
    # Example usage (integrate with your existing code)
    create_full_page_qr_doc(output_dir="data/processed", docx_path="qr_codes_landscape_a4.docx")
    